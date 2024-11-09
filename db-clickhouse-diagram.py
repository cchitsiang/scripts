"""
ClickHouse Data Dictionary Generator

This script connects to a ClickHouse database and generates a comprehensive data dictionary
in Microsoft Word format. It performs the following tasks:
- Connects to a local ClickHouse database
- Retrieves metadata about all tables and columns in the current database
- Processes special data types like Enums and handles foreign key mappings
- Generates a formatted Word document containing:
  * Table listings
  * Column details (name, data type, nullability, and foreign key references)
  * Formatted tables with consistent styling
  
Output: Creates a Word document named 'clickhouse_data_dictionary.docx'
"""

from clickhouse_driver import Client
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Connect to ClickHouse
print("Attempting database connection...")
client = Client(
    host='localhost',
    port=8811,
    database='default',
    user='default',
    password='',
)
print("Database connection successful!")

heading = "ClickHouse Data Dictionary"
filename = "clickhouse_data_dictionary.docx"
# Add this dictionary for foreign key mappings
FK_MAPPINGS = {
    'stepId': 'postgres.step.id',
    'customerId': 'mongo.customers.id',
    'templateId': 'postgres.template.id',
    'campaignId': 'postgres.campaign.id'
}

# Create a Word document
doc = Document()
doc.add_heading(heading, level=1)

# Function to fetch table metadata
def fetch_table_metadata():
    print("Fetching table metadata...")
    
    # Updated query for ClickHouse columns
    columns_query = """
    SELECT 
        table,
        name as column_name,
        type as data_type,
        default_kind
    FROM system.columns
    WHERE database = currentDatabase()
    ORDER BY table, position
    """
    
    # Fetch columns
    print("Fetching columns...")
    columns = client.execute(columns_query)
    
    # Create empty dictionaries
    fk_dict = {}  # We won't use this for ClickHouse
    enum_dict = {}  # We'll handle enums directly from the type string
    
    print(f"Found {len(columns)} columns")
    return columns, fk_dict, enum_dict

# Organize data by table
print("Organizing table data...")
tables = {}
columns, fk_dict, enum_dict = fetch_table_metadata()
for row in columns:
    table_name, column_name, data_type, default_kind = row
    
    # Handle data type conversion
    final_data_type = data_type
    
    # Handle Enum types if they exist in the data type string
    if 'Enum' in data_type:
        try:
            # Extract enum values from the data type string
            enum_values = data_type.split('(')[1].split(')')[0]
            values = [v.split("'")[1] for v in enum_values.split(',')]
            final_data_type = f"enum({', '.join(values)})"
        except:
            final_data_type = data_type
    elif 'String' in data_type:
        final_data_type = 'string'
    
    # Get foreign key information from hardcoded mappings
    remarks = ""
    if column_name in FK_MAPPINGS:
        remarks = f"References {FK_MAPPINGS[column_name]}"
    
    if table_name not in tables:
        tables[table_name] = []
        print(f"Processing table: {table_name}")
    
    tables[table_name].append({
        "column_name": column_name,
        "data_type": final_data_type,
        "is_nullable": "Yes" if "Nullable" in data_type else "No",
        "remarks": remarks
    })

# Write table information to Word
print("Writing to Word document...")
for table, columns in tables.items():
    doc.add_heading(f"Table: {table}", level=2)
    table_doc = doc.add_table(rows=1, cols=4)
    
    # Set table style
    table_doc.style = 'Table Grid'
    table_doc.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Update header row
    hdr_cells = table_doc.rows[0].cells
    hdr_cells[0].text = "Column Name"
    hdr_cells[1].text = "Data Type"
    hdr_cells[2].text = "Nullable"
    hdr_cells[3].text = "Remarks"

    # Style header row
    for cell in table_doc.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="DDDDDD"/>'))
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if not run:
            run = paragraph.add_run()
        else:
            run = run[0]
        run.font.bold = True

    # Add data rows
    for column in columns:
        row_cells = table_doc.add_row().cells
        row_cells[0].text = column["column_name"]
        row_cells[1].text = column["data_type"]
        row_cells[2].text = column["is_nullable"]
        row_cells[3].text = column["remarks"]

    # Add spacing after table
    doc.add_paragraph()

# Save the document
print("Saving document...")
doc.save(filename)
print("Document saved successfully!")
