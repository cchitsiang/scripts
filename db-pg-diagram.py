import psycopg2
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Connect to PostgreSQL
print("Attempting database connection...")
conn = psycopg2.connect(
    dbname="postgres",
    user="postgres",
    password="password",
    host="localhost",
    port="5432"
)
print("Database connection successful!")

heading = "Postgres Data Dictionary"
filename = "postgres_data_dictionary.docx"

# Create a Word document
doc = Document()
doc.add_heading(heading, level=1)

# Function to fetch table metadata
def fetch_table_metadata():
    print("Fetching table metadata...")
    
    # Query to get enum types and their values
    enum_query = """
    SELECT 
        t.typname as enum_name,
        e.enumlabel as enum_value
    FROM pg_type t 
    JOIN pg_enum e ON t.oid = e.enumtypid
    JOIN pg_catalog.pg_namespace n ON n.oid = t.typnamespace
    WHERE n.nspname = 'public'
    ORDER BY t.typname, e.enumsortorder;
    """
    
    columns_query = """
    SELECT 
        table_name, 
        column_name, 
        data_type,
        udt_name,  -- Add this to get the actual enum type name
        is_nullable
    FROM information_schema.columns
    WHERE table_schema = 'public'
    ORDER BY table_name, ordinal_position;
    """
    
    fk_query = """
    SELECT
        tc.table_name, 
        kcu.column_name,
        ccu.table_name AS foreign_table_name,
        ccu.column_name AS foreign_column_name
    FROM information_schema.table_constraints AS tc 
    JOIN information_schema.key_column_usage AS kcu
        ON tc.constraint_name = kcu.constraint_name
        AND tc.table_schema = kcu.table_schema
    JOIN information_schema.constraint_column_usage AS ccu
        ON ccu.constraint_name = tc.constraint_name
        AND ccu.table_schema = tc.table_schema
    WHERE tc.constraint_type = 'FOREIGN KEY'
    AND tc.table_schema = 'public';
    """
    
    with conn.cursor() as cur:
        # Fetch enum types and their values
        cur.execute(enum_query)
        enums = cur.fetchall()
        
        # Create enum dictionary
        enum_dict = {}
        for enum_name, enum_value in enums:
            if enum_name not in enum_dict:
                enum_dict[enum_name] = []
            enum_dict[enum_name].append(enum_value)
        
        # Fetch columns
        cur.execute(columns_query)
        columns = cur.fetchall()
        
        # Fetch foreign keys
        cur.execute(fk_query)
        fks = cur.fetchall()
        
        # Create FK lookup dictionary
        fk_dict = {}
        for fk in fks:
            table_name, column_name, foreign_table, foreign_column = fk
            key = f"{table_name}.{column_name}"
            fk_dict[key] = f"References {foreign_table}({foreign_column})"
            
        print(f"Found {len(columns)} columns, {len(fks)} foreign keys, and {len(enum_dict)} enum types")
        return columns, fk_dict, enum_dict

# Organize data by table
print("Organizing table data...")
tables = {}
columns, fk_dict, enum_dict = fetch_table_metadata()
for row in columns:
    table_name, column_name, data_type, udt_name, is_nullable = row
    
    # Handle data type conversion
    if data_type == 'character varying':
        final_data_type = 'string'
    elif data_type == 'USER-DEFINED':
        # Check if it's an enum type
        if udt_name in enum_dict:
            enum_values = ', '.join(enum_dict[udt_name])
            final_data_type = f"enum({enum_values})"
        else:
            final_data_type = data_type
    else:
        final_data_type = data_type
    
    # Get foreign key information if exists
    fk_info = fk_dict.get(f"{table_name}.{column_name}", "")
    
    if table_name not in tables:
        tables[table_name] = []
        print(f"Processing table: {table_name}")
    
    tables[table_name].append({
        "column_name": column_name,
        "data_type": final_data_type,
        "is_nullable": is_nullable,
        "remarks": fk_info
    })

# Write table information to Word
print("Writing to Word document...")
for table, columns in tables.items():
    doc.add_heading(f"Table: {table}", level=2)
    table_doc = doc.add_table(rows=1, cols=4)
    
    # Set table style
    table_doc.style = 'Table Grid'
    table_doc.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Update header row
    hdr_cells = table_doc.rows[0].cells
    hdr_cells[0].text = "Column Name"
    hdr_cells[1].text = "Data Type"
    hdr_cells[2].text = "Nullable"
    hdr_cells[3].text = "Remarks"

    # Style header row
    for cell in table_doc.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="DDDDDD"/>'))
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if not run:
            run = paragraph.add_run()
        else:
            run = run[0]
        run.font.bold = True

    # Add data rows
    for column in columns:
        row_cells = table_doc.add_row().cells
        row_cells[0].text = column["column_name"]
        row_cells[1].text = column["data_type"]
        row_cells[2].text = column["is_nullable"]
        row_cells[3].text = column["remarks"]

    # Add spacing after table
    doc.add_paragraph()

# Save the document
doc.save(filename)

# Close connection
conn.close()
